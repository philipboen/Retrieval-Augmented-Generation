from .base_parser import BaseParser
import logging
import io
import pytesseract
from PIL import Image
import docx
from .pdf_parser import PdfParser
from docx.document import Document

class DocxParser(BaseParser):
    def __init__(self):
        if not PdfParser._verify_tesseract():
            raise RuntimeError("Tesseract OCR is not properly installed")

    def parse(self, filepath: str) -> str:
        try:
            doc = docx.Document(filepath)
            content = []

            # Extract content in order: TOC, paragraphs, tables, images
            content.extend(self._extract_toc(doc))
            content.extend(self._extract_paragraphs(doc))
            content.extend(self._extract_tables(doc))
            content.extend(self._extract_images(doc))

            final_content = " ".join(filter(None, content))
            if not final_content.strip():
                raise ValueError("No content extracted from DOCX")

            return final_content
        except Exception as e:
            logging.error(f"Error processing DOCX: {e}")
            return f"Error processing DOCX file: {str(e)}"

    def _extract_toc(self, doc: Document) -> list:
        """Extract Table of Contents"""
        toc_content = []
        try:
            for para in doc.paragraphs:
                if para.style.name.lower().startswith(("toc", "table of contents")):
                    text = para.text.strip()
                    if text:
                        toc_content.append(text)
            logging.info(f"Extracted {len(toc_content)} TOC entries")
        except Exception as e:
            logging.error(f"TOC extraction failed: {e}")
        return toc_content

    def _extract_paragraphs(self, doc: Document) -> list:
        """Extract paragraph content"""
        para_content = []
        try:
            for para in doc.paragraphs:
                if not para.style.name.lower().startswith(("toc", "table of contents")):
                    text = para.text.strip()
                    if text:
                        para_content.append(text)
            logging.info(f"Extracted {len(para_content)} paragraphs")
        except Exception as e:
            logging.error(f"Paragraph extraction failed: {e}")
        return para_content

    def _extract_tables(self, doc: Document) -> list:
        """Extract table content"""
        table_content = []
        try:
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    ]
                    if cells:
                        rows.append(" | ".join(cells))
                if rows:
                    table_content.append("\n".join(rows))
            logging.info(f"Extracted content from {len(doc.tables)} tables")
        except Exception as e:
            logging.error(f"Table extraction failed: {e}")
        return table_content

    def _extract_images(self, doc: Document) -> list:
        """Extract and process images with OCR"""
        image_content = []
        try:
            for shape in doc.inline_shapes:
                if shape.type == 3:  # Picture type
                    try:
                        image_bytes = (
                            shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                        )
                        if image_bytes:
                            image = Image.open(io.BytesIO(image_bytes))
                            if image.mode != "RGB":
                                image = image.convert("RGB")
                            ocr_text = pytesseract.image_to_string(
                                image, config="--psm 3 --oem 3"
                            ).strip()
                            if ocr_text:
                                image_content.append(ocr_text)
                    except Exception as img_e:
                        logging.error(f"Failed to process image: {img_e}")
            logging.info(f"Processed {len(image_content)} images with OCR")
        except Exception as e:
            logging.error(f"Image extraction failed: {e}")
        return image_content
